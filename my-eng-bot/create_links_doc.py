import sys
import subprocess

try:
    import docx
    from docx.shared import Pt
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Pt

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'

add_heading('Ссылки и материалы для самостоятельного изучения (GWARD)', 0)

doc.add_paragraph('В этом документе собраны прямые ссылки на статьи, аналитику и видео, которые помогут вам глубже погрузиться в контекст перед собеседованием.')

# Блок 1
add_heading('1. Главное интервью Генерального директора (Александр Ермаченко)', 1)
doc.add_paragraph('Это базовый текст, на котором строится большинство выводов о стратегии компании. В нём Ермаченко рассказывает историю создания, почему они ушли от простой торговли к производству, и как они решают проблемы клиентов.')
doc.add_paragraph('Ссылка на статью "Бизнес в перчатках":')
doc.add_paragraph('https://getsiz.ru/biznes-v-perchatkakh.html')
doc.add_paragraph('На что обратить внимание при чтении:')
doc.add_paragraph('- Как он описывает боль с поставщиками до 2012 года (почему решили делать всё сами).')
doc.add_paragraph('- Кейс про сварщиков (влагостойкие краги) и сборщиков авиадвигателей (противоударные перчатки). Это показывает фокус компании на R&D.')
doc.add_paragraph('- Заявление про "складские запасы на 6-7 месяцев и отгрузку за 48 часов" — это та самая модель, которую вам, возможно, предстоит менять как Комдиру.')

# Блок 2
add_heading('2. Аналитика рынка СИЗ в России (Тренды)', 1)
doc.add_paragraph('Чтобы говорить с собственниками на одном языке, нужно понимать макроэкономику рынка. Портал Гетсиз (Getsiz.ru) — это главное отраслевое СМИ в России.')
doc.add_paragraph('Статья "11 трендов российского рынка СИЗ в 2026 году":')
doc.add_paragraph('https://getsiz.ru/11-trendov-rossiiskogo-rynka-siz-v-2026-godu.html')
doc.add_paragraph('На что обратить внимание при чтении:')
doc.add_paragraph('- Тренд на рост онлайн-торговли СИЗ (B2B маркетплейсы вроде Комус и ВсеИнструменты).')
doc.add_paragraph('- Влияние Постановления Правительства № 1875 (национальный режим закупок — импортозамещение).')
doc.add_paragraph('- Возвращение зарубежных производителей (китайская экспансия и параллельный импорт).')

# Блок 3
add_heading('3. Изучение продукта: Официальные ресурсы Gward', 1)
doc.add_paragraph('Обязательно полистайте их сайт, чтобы визуально запомнить хиты продаж и навигацию. Это покажет HR вашу вовлеченность.')
doc.add_paragraph('Официальный сайт (раздел О компании):')
doc.add_paragraph('https://gward.ru/about/')
doc.add_paragraph('Каталог продукции (обратите внимание на разделы "Спилковые краги" и "С нитриловым покрытием"):')
doc.add_paragraph('https://gward.ru/catalog/')

# Блок 4
add_heading('4. Что почитать про ваших конкурентов', 1)
doc.add_paragraph('Вы знаете рынок по Ansell и Portwest, но полезно освежить знания о текущих российских гигантах, с которыми бьется Gward.')
doc.add_paragraph('Рейтинг крупнейших компаний рынка СИЗ (полезно пробежаться глазами по Топ-10, чтобы понимать вес "Восток-Сервис", "Техноавиа" и "Тракт"):')
doc.add_paragraph('https://getsiz.ru/60-krupnejshih-kompanij-rossijskogo-rynka-specodezhdy-i-siz-2025-goda.html')

doc.save('GWARD_Reading_Links.docx')
print("Links document generated successfully.")
