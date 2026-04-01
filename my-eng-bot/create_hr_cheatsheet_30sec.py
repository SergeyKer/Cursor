import sys
import subprocess

try:
    import docx
    from docx.shared import Pt
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Pt


doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)


def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'


add_heading('HR Cheatsheet 30 sec (GWARD)', 0)
doc.add_paragraph('7 фраз, которые нужно держать в голове на звонке/встрече.')

phrases = [
    '1) «Я коммерческий руководитель с 20+ годами в B2B-дистрибуции и управлении P&L.»',
    '2) «Мой самый релевантный опыт — рынок СИЗ рук: Ansell (РФ/СНГ) и Portwest (дистрибуция в ЦФО).»',
    '3) «Я не про ручные продажи, я про систему: CRM, KPI, маржинальность и финансовая дисциплина.»',
    '4) «В Blesk InCare я успешно вел коммерцию, а затем усилил кросс-функциональные процессы в сервисе и логистике.»',
    '5) «Сейчас осознанно фокусируюсь на своей сильной стороне — роли коммерческого директора в B2B.»',
    '6) «В Gward мне интересен масштаб, продукт и задача кратного роста через системную коммерческую модель.»',
    '7) «Если нужно, отдельно раскрою мой опыт в СИЗ и работе с дилерской сетью — это максимально релевантно вашей вакансии.»',
]

for line in phrases:
    doc.add_paragraph(line)

add_heading('Короткий ответ про уход (если спросят в лоб)', 1)
doc.add_paragraph(
    '«Это был переход в сложный кросс-функциональный проект. После закрытия ключевых задач '
    'я осознанно вернулся к своей основной экспертизе — B2B-коммерции, росту и P&L».'
)

doc.save('HR_cheatsheet_30sec.docx')
print('Created HR_cheatsheet_30sec.docx')
